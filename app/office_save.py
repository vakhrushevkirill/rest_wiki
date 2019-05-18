from openpyxl import Workbook, load_workbook
from docx import Document
import os
path_temp = 'temp/'


class Excel():
    def __init__(self, title='test'):
        """Создает файл excel
        :param title: имя файла"""
        wb = Workbook()
        ws = wb.active
        ws.title = "JSON лист"
        self.title = title + '.xlsx'
        self.filename = path_temp + title + '.xlsx'
        wb.save(self.filename)
    
    def create(self, data):
        
        wb = load_workbook(self.filename)
        ws = wb.active
        row = 1
        column = 1
        if isinstance(data, dict):
            row, column = self.is_dict(data, ws, row, column)
        else:
            row, column = self.is_list(data, ws, row, column)
        wb.save(self.filename)
        wb.close()

    def is_list(self, list_data, ws, row, column):
        for item in list_data:
            if isinstance(item, dict):
                row, column = self.is_dict(item, ws, row, column)
            elif isinstance(item, list):
                row, column = self.is_list(item, ws, row, column)
            else:
                ws.cell(row, column, value=item)
                row += 1
        return row, column

    def is_dict(self, dict_data, ws, row, column):
        for key in dict_data:
            if isinstance(dict_data[key], dict):
                ws.cell(row, column, value=key)
                row, _ = self.is_dict(dict_data[key], ws, row, column+1)
            elif isinstance(dict_data[key], list):
                ws.cell(row, column, value=key)
                row, _ = self.is_list(dict_data[key], ws, row, column+1)
            else: 
                ws.cell(row, column, value=key)
                column += 1
                ws.cell(row, column, value=str(dict_data[key]))
                row += 1
                column -= 1
        return row, column

    def __del__(self):
        os.remove(self.filename)



class Word():
    def __init__(self, title='test'):
        self.document = Document()
        self.title = title + '.doc'
        self.filename = path_temp + title + '.doc'
        self.spaces_count = 4

    def create(self, data):
        spaces = 0
        if isinstance(data, dict):
            spaces = self.is_dict(data, spaces)
        else:
            spaces = self.is_list(data, spaces)
        self.document.save(self.filename)

    def is_list(self, list_data, spaces):
        for item in list_data:
            if isinstance(item, dict):
                spaces = self.is_dict(item, spaces)
            elif isinstance(item, list):
                spaces = self.is_list(item, spaces)
            else:
                self.document.add_paragraph(' ' * self.spaces_count + item)
                spaces += 4
        return spaces

    def is_dict(self, dict_data, spaces):
        for key in dict_data:
            if isinstance(dict_data[key], dict):
                self.document.add_paragraph(' ' * self.spaces_count + key)
                spaces = self.is_dict(dict_data[key], spaces+4)
            elif isinstance(dict_data[key], list):
                self.document.add_paragraph(' ' * self.spaces_count + key)
                spaces = self.is_list(dict_data[key], spaces+4)
            else: 
                self.document.add_paragraph(
                   str(key) + ': ' + str(dict_data[key])
                   )
        return spaces

    def __del__(self):
        os.remove(self.filename)


if __name__ == "__main__":
    def test_excel():
        from rest import RequestToWiki, RestApiWiki
        ex = Excel()
        rq_to_wiki = RequestToWiki()
        rq_to_wiki.get_level_list('Sad')
        print(rq_to_wiki.list_level)
        ex.create(rq_to_wiki.list_level)

    def test_doc():
        from rest import RequestToWiki, RestApiWiki
        doc = Word()
        rq_to_wiki = RequestToWiki()
        rq_to_wiki.get_level_list('Sad')
        doc.create(rq_to_wiki.list_level)

    test_doc()